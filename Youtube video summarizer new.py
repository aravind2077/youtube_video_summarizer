from transformers import pipeline
from youtube_transcript_api import YouTubeTranscriptApi
import re
import pickle
import docx

def get_id(youtube_video):
    video_id= re.findall(r'watch\?v=(.+)', youtube_video)
    return video_id

def get_rawbody(transcript):
    body = ""
    for i in transcript:
        body += ' ' + i['text']
        
    return body

def get_summary(body, partsumlen):
    filename = 'finalized_model.sav'
    summarizer = pickle.load(open(filename, 'rb'))
    summarized_text = []
    num_iters = int(len(body)/500)
    for i in range(0, num_iters + 1):
        start = 0
        partresult_text=''
        partresult_list=[]
        start = i * 500
        end = (i + 1) * 500
        partresult_list= body[start:end]
        for j in partresult_list:
            partresult_text+= ' ' + j
        out = summarizer(partresult_text, min_length= partsumlen, max_length= partsumlen+50)
        out = out[0]
        out = out['summary_text']
        summarized_text.append(out)
        
    return summarized_text

video_id= get_id(str(input('Enter Youtube URL: ')))[0]

try:
    YouTubeTranscriptApi.get_transcript(video_id)
    transcript = YouTubeTranscriptApi.get_transcript(video_id)
    status= True
except:
    print("No transcript available")
    status= False
    
if status == True:
    body= get_rawbody(transcript)
    
    transc= docx.Document()
    transc.add_paragraph(body)
    transc.save('Transcript.docx')
    print('Transcript is saved in Transcript.docx')
    
    body= body.split()
    videolen= len(body)
    stat= False
    while stat != True:
        print(f"The video's transcription is {videolen} words long")
        try:
            sumlen= int(input('How many words would you like the Summary to be?'))
            stat= True
            if sumlen > videolen:
                print("The Summary can't be bigger than the Transcription.\nPlease try a lower value")
                stat= False
        except:
            print('Please enter a Natural Number')
            stat= False
    
    partsumlen= int(sumlen/((videolen/500)+1))

    result= get_summary(body, partsumlen)
    doc= docx.Document()
    doc.add_paragraph(result)
    doc.save('Summary.docx')
    print('Summarised text is saved in Summary.docx')
    
else:
    print("bye")